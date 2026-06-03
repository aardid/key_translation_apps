"""
Extract parallel EN-ES translation data from two Word documents.

Reads the English original and Spanish translation of SOP_TY_001,
aligns paragraphs and table cells, and outputs a JSON file with
paired translations.

Strategy:
  1. Walk each document body XML to get paragraphs and tables in
     document order.
  2. Identify heading paragraphs and match them between EN and ES
     using a known translation map plus fuzzy position-based fallback.
  3. Between matched heading anchors, pair non-empty paragraphs 1-to-1
     and pair tables 1-to-1 for cell extraction.
  4. For tables, pair cells by (row, col) position, deduplicating
     merged cells.
"""

import json
import os
import re
import docx
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Known heading translations (EN -> ES) for reliable anchoring
# ---------------------------------------------------------------------------
HEADING_MAP = {
    "PURPOSE": "PROPOSITO",
    "SCOPE": "ALCANCE",
    "RESPONSIBILITY": "RESPONSABILIDAD",
    "MAINTENANCE": "MANTENIMIENTO",
    "DEFINITIONS": "DEFINICIONES",
    "RISK ASSESSMENT": "EVALUACION DE RIESGOS",
    "EMERGENCY RESPONSE PLANNING": "PLANIFICACION DE RESPUESTA ANTE EMERGENCIAS",
    "EQUIPMENT SAFETY": "SEGURIDAD DE LOS EQUIPOS",
    "INCIDENT INVESTIGATION AND REPORTING PROCEDURE": "PROCEDIMIENTO DE INVESTIGACION Y REPORTE DE INCIDENTES",
    "GENERAL HEALTH AND SAFETY": "SALUD Y SEGURIDAD GENERAL",
    "EQUIPMENT INSPECTION AND MAINTENANCE SCHEDULE": "PROGRAMA DE INSPECCION Y MANTENCION DE EQUIPOS",
    "RECORDKEEPING": "MANTENCION DE REGISTROS",
    "REFERENCES & RELATED DOCUMENTS": "REFERENCIAS Y DOCUMENTOS RELACIONADOS",
    "APPENDIX A: WORK PROCEDURES": "ANEXO A: PROCEDIMIENTOS DE TRABAJO",
    "APPENDIX B:": "ANEXO B:",
    "APPENDIX C:": "ANEXO C:",
    "APPENDIX D:": "ANEXO D:",
    "APPENDIX F: TYPHOON COMPONENTS": "ANEXO F: COMPONENTES DEL TYPHOON",
    # Heading 2 level
    "Zero Harm is achievable.": "Se puede alcanzar Cero Dano.",
    "About Typhoon": "Acerca de Typhoon",
    "Significant hazards": "Peligros significativos",
    "Emergency contacts": "Contactos de emergencia",
    "Emergency procedures": "Procedimientos de emergencia",
    "Personnel tracking": "Control del personal",
    "Emergency Shutdown": "Apagado de emergencia",
    "Emergency Response Drills": "Simulacros de respuesta ante emergencias",
    "Health & Safety signs": "Senalizacion de salud y seguridad",
    "Isolation procedures": "Procedimientos de aislamiento",
    "Working at heights": "Trabajo en altura",
    "Hazardous material handling": "Manipulacion de materiales peligrosos",
    "EMERGENCY RESPONSE PLAN SAMPLE TEMPLATES": "PLANTILLAS DE EJEMPLO DEL PLAN DE RESPUESTA A EMERGENCIAS",
    "TYPHOON GEOPHYSICS TRANSMITTER CHECKLIST": "CHECKLIST DEL TRANSMISOR GEOFISICO TYPHOON",
    "TRANSMISSION PHASE DAILY CHECKLIST": "CHECKLIST DIARIO DE LA FASE DE TRANSMISION",
    "HIGH POWER EQUIPMENT DAILY SITE AND EQUIPMENT INSPECTION FORM": "FORMULARIO DE INSPECCION DIARIA EN TERRENO",
    "Selecting the correct wire": "Seleccion del cable adecuado",
    "Multimeter": "Multimetro",
    "Typhoon beacon": "Baliza del Typhoon",
    "Typhoon container": "Contenedor del Typhoon",
    "Grounding and Current Electrodes pits": "Puesta a tierra y pozos de electrodos de corriente",
    "Low voltage line test": "Prueba de linea de baja tension",
    "Open Output Error": "Error Open Output",
    "Open circuit in transmission line": "Circuito abierto en la linea de transmision",
}


def normalize_for_match(text):
    """Strip accents, special chars, collapse whitespace for fuzzy matching."""
    if not text:
        return ""
    # Replace non-breaking spaces
    text = text.replace("\xa0", " ")
    # Remove accents (simple approach for Spanish)
    replacements = {
        "\xe1": "a", "\xe9": "e", "\xed": "i", "\xf3": "o", "\xfa": "u",
        "\xc1": "A", "\xc9": "E", "\xcd": "I", "\xd3": "O", "\xda": "U",
        "\xf1": "n", "\xd1": "N",
        "\xfc": "u", "\xdc": "U",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Collapse whitespace, uppercase
    text = " ".join(text.split()).upper().strip()
    return text


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def clean_text(text):
    """Normalise whitespace in extracted text."""
    if text is None:
        return ""
    text = text.replace("\xa0", " ")
    text = " ".join(text.split())
    return text.strip()


def is_nonempty(text):
    """Return True if text has actual content."""
    return bool(clean_text(text))


# ---------------------------------------------------------------------------
# Document parsing
# ---------------------------------------------------------------------------

def parse_document(doc):
    """Walk the body XML and return an ordered list of elements."""
    p_idx = 0
    t_idx = 0
    elements = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            if p_idx < len(doc.paragraphs):
                para = doc.paragraphs[p_idx]
                elements.append({
                    "type": "paragraph",
                    "index": p_idx,
                    "text": para.text.strip(),
                    "style": para.style.name,
                })
            p_idx += 1
        elif child.tag == qn("w:tbl"):
            elements.append({
                "type": "table",
                "index": t_idx,
            })
            t_idx += 1
    return elements


def is_heading(element):
    """Check if an element is a non-empty heading paragraph."""
    if element["type"] != "paragraph":
        return False
    if not element["style"].startswith("Heading"):
        return False
    return is_nonempty(element["text"])


def split_by_headings(elements):
    """Split elements into sections delimited by heading paragraphs.

    Returns a list of (heading_element_or_None, [body_elements]).
    """
    sections = []
    current_heading = None
    current_body = []

    for elem in elements:
        if is_heading(elem):
            sections.append((current_heading, current_body))
            current_heading = elem
            current_body = []
        else:
            current_body.append(elem)

    sections.append((current_heading, current_body))
    return sections


def match_headings(en_sections, es_sections):
    """Match EN and ES sections by heading content using the translation map.

    Returns list of (en_section_index, es_section_index) pairs.
    """
    # Build normalized heading map
    norm_map = {}
    for en_key, es_val in HEADING_MAP.items():
        en_norm = normalize_for_match(en_key)
        es_norm = normalize_for_match(es_val)
        norm_map[en_norm] = es_norm

    # Get normalized heading texts
    en_headings = []
    for i, (heading, _) in enumerate(en_sections):
        if heading is not None:
            en_headings.append((i, normalize_for_match(heading["text"])))

    es_headings = []
    for i, (heading, _) in enumerate(es_sections):
        if heading is not None:
            es_headings.append((i, normalize_for_match(heading["text"])))

    # Match by translation map
    pairs = []
    used_es = set()

    # First pass: exact map matches
    for en_idx, en_norm in en_headings:
        # Check if en_norm matches a key in our map
        expected_es = None
        for map_en, map_es in norm_map.items():
            if map_en in en_norm or en_norm in map_en:
                expected_es = map_es
                break

        if expected_es:
            # Find the matching ES heading
            for es_idx, es_norm in es_headings:
                if es_idx not in used_es:
                    if expected_es in es_norm or es_norm in expected_es:
                        pairs.append((en_idx, es_idx))
                        used_es.add(es_idx)
                        break

    # Second pass: try positional matching for any unmatched headings
    matched_en = {p[0] for p in pairs}
    unmatched_en = [(i, h) for i, h in en_headings if i not in matched_en]
    unmatched_es = [(i, h) for i, h in es_headings if i not in used_es]

    # Match remaining by position (greedy)
    for en_idx, en_norm in unmatched_en:
        best_es = None
        best_dist = float("inf")
        for es_idx, es_norm in unmatched_es:
            if es_idx not in used_es:
                # Use section index distance as heuristic
                dist = abs(en_idx - es_idx)
                if dist < best_dist:
                    best_dist = dist
                    best_es = es_idx

        if best_es is not None and best_dist <= 3:
            pairs.append((en_idx, best_es))
            used_es.add(best_es)

    # Sort by EN index
    pairs.sort()

    # Also include pre-heading sections (index 0)
    if en_sections[0][0] is None and es_sections[0][0] is None:
        pairs.insert(0, (0, 0))

    return pairs


# ---------------------------------------------------------------------------
# Table cell extraction
# ---------------------------------------------------------------------------

def extract_table_cells(en_table, es_table):
    """Extract aligned cell texts from two tables."""
    results = []
    rows = min(len(en_table.rows), len(es_table.rows))
    cols = min(len(en_table.columns), len(es_table.columns))

    seen = set()
    for r in range(rows):
        for c in range(cols):
            en_text = clean_text(en_table.cell(r, c).text)
            es_text = clean_text(es_table.cell(r, c).text)

            if not en_text or not es_text:
                continue

            en_cell_id = id(en_table.cell(r, c)._element)
            es_cell_id = id(es_table.cell(r, c)._element)
            cell_key = (en_cell_id, es_cell_id)
            if cell_key in seen:
                continue
            seen.add(cell_key)

            results.append({
                "en": en_text,
                "es": es_text,
                "source": "SOP_TY_001",
                "type": "table_cell",
            })

    return results


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base_dir = r"U:\Research\EruptionForecasting\fires\Wintel\key_translation"
    example_dir = os.path.join(base_dir, "translation_example")

    en_path = os.path.join(
        example_dir,
        "SOP_TY_001 Typhoon Safe Operating Procedure.docx",
    )
    es_path = os.path.join(
        example_dir,
        "Traduccion SOP_TY_001 Typhoon Safe Operating Procedure.docx",
    )

    print("Loading English document...")
    en_doc = docx.Document(en_path)
    print("Loading Spanish document...")
    es_doc = docx.Document(es_path)

    print(f"EN: {len(en_doc.paragraphs)} paragraphs, {len(en_doc.tables)} tables")
    print(f"ES: {len(es_doc.paragraphs)} paragraphs, {len(es_doc.tables)} tables")

    # Parse body elements
    en_elements = parse_document(en_doc)
    es_elements = parse_document(es_doc)

    # Split by headings
    en_sections = split_by_headings(en_elements)
    es_sections = split_by_headings(es_elements)

    print(f"EN heading-delimited sections: {len(en_sections)}")
    print(f"ES heading-delimited sections: {len(es_sections)}")

    # Match headings
    section_pairs = match_headings(en_sections, es_sections)
    print(f"Matched section pairs: {len(section_pairs)}")

    # Show matched headings for verification
    print("\n--- Heading alignment ---")
    for en_idx, es_idx in section_pairs:
        en_h = en_sections[en_idx][0]
        es_h = es_sections[es_idx][0]
        en_txt = clean_text(en_h["text"])[:45] if en_h else "(pre-heading)"
        es_txt = clean_text(es_h["text"])[:45] if es_h else "(pre-heading)"
        print(f"  EN[{en_idx:2d}] {en_txt:45s} <-> ES[{es_idx:2d}] {es_txt}")

    results = []
    stats = {
        "heading_pairs": 0,
        "paragraph_pairs": 0,
        "table_cell_pairs": 0,
        "skipped_en_paras": 0,
        "skipped_es_paras": 0,
        "skipped_en_tables": 0,
        "skipped_es_tables": 0,
    }

    for en_idx, es_idx in section_pairs:
        en_heading, en_body = en_sections[en_idx]
        es_heading, es_body = es_sections[es_idx]

        # Pair the heading text
        if en_heading is not None and es_heading is not None:
            en_h_text = clean_text(en_heading["text"])
            es_h_text = clean_text(es_heading["text"])
            if en_h_text and es_h_text:
                results.append({
                    "en": en_h_text,
                    "es": es_h_text,
                    "source": "SOP_TY_001",
                    "type": "paragraph",
                })
                stats["heading_pairs"] += 1

        # Separate body elements into non-empty paragraphs and tables
        en_paras = [
            e for e in en_body
            if e["type"] == "paragraph" and is_nonempty(e["text"])
        ]
        es_paras = [
            e for e in es_body
            if e["type"] == "paragraph" and is_nonempty(e["text"])
        ]
        en_tables = [e for e in en_body if e["type"] == "table"]
        es_tables = [e for e in es_body if e["type"] == "table"]

        # Pair paragraphs 1-to-1
        n_paras = min(len(en_paras), len(es_paras))
        stats["skipped_en_paras"] += max(0, len(en_paras) - n_paras)
        stats["skipped_es_paras"] += max(0, len(es_paras) - n_paras)

        for pi in range(n_paras):
            en_text = clean_text(en_paras[pi]["text"])
            es_text = clean_text(es_paras[pi]["text"])
            if en_text and es_text:
                results.append({
                    "en": en_text,
                    "es": es_text,
                    "source": "SOP_TY_001",
                    "type": "paragraph",
                })
                stats["paragraph_pairs"] += 1

        # Pair tables 1-to-1
        n_tables = min(len(en_tables), len(es_tables))
        stats["skipped_en_tables"] += max(0, len(en_tables) - n_tables)
        stats["skipped_es_tables"] += max(0, len(es_tables) - n_tables)

        for ti in range(n_tables):
            en_tidx = en_tables[ti]["index"]
            es_tidx = es_tables[ti]["index"]

            if en_tidx < len(en_doc.tables) and es_tidx < len(es_doc.tables):
                cells = extract_table_cells(
                    en_doc.tables[en_tidx],
                    es_doc.tables[es_tidx],
                )
                results.extend(cells)
                stats["table_cell_pairs"] += len(cells)

    print(f"\nExtraction complete:")
    print(f"  Heading pairs: {stats['heading_pairs']}")
    print(f"  Body paragraph pairs: {stats['paragraph_pairs']}")
    print(f"  Table cell pairs: {stats['table_cell_pairs']}")
    print(f"  Total pairs: {len(results)}")
    print(f"  Skipped EN paragraphs (no ES match): {stats['skipped_en_paras']}")
    print(f"  Skipped ES paragraphs (no EN match): {stats['skipped_es_paras']}")
    print(f"  Skipped EN tables (no ES match): {stats['skipped_en_tables']}")
    print(f"  Skipped ES tables (no EN match): {stats['skipped_es_tables']}")

    # Write output
    output_path = os.path.join(base_dir, "translations_extracted.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nOutput written to: {output_path}")

    # Show sample entries
    print("\n--- Sample paragraph entries ---")
    para_entries = [e for e in results if e["type"] == "paragraph"]
    if para_entries:
        indices = [0, 1, 5, 10, len(para_entries) // 4,
                   len(para_entries) // 2, 3 * len(para_entries) // 4,
                   -2, -1]
        shown = set()
        for idx in indices:
            actual = idx if idx >= 0 else len(para_entries) + idx
            if actual not in shown and 0 <= actual < len(para_entries):
                shown.add(actual)
                e = para_entries[actual]
                print(f"  [{actual}] EN: {e['en'][:90]}")
                print(f"       ES: {e['es'][:90]}")
                print()

    print("--- Sample table cell entries ---")
    table_entries = [e for e in results if e["type"] == "table_cell"]
    if table_entries:
        indices = [0, 2, len(table_entries) // 3,
                   2 * len(table_entries) // 3, -1]
        shown = set()
        for idx in indices:
            actual = idx if idx >= 0 else len(table_entries) + idx
            if actual not in shown and 0 <= actual < len(table_entries):
                shown.add(actual)
                e = table_entries[actual]
                print(f"  [{actual}] EN: {e['en'][:90]}")
                print(f"       ES: {e['es'][:90]}")
                print()


if __name__ == "__main__":
    main()
